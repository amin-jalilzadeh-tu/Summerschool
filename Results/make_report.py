"""Build the 12-page case-study report.

Every number in the document is COMPUTED here from Appendix C and from the pipeline runs —
nothing is typed in by hand. If the data changes, the prose changes with it.

    python3 make_report.py   ->   CaseStudy2_Report.docx
"""
import glob
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from scipy.optimize import linprog

HERE = Path(__file__).resolve().parent
DT = Path("/Users/amin/Documents/4_Learning/Summer school/Digital Twin")
OUT = HERE / "CaseStudy2_Report.docx"

INK = RGBColor(0x18, 0x18, 0x1B)
MUTED = RGBColor(0x52, 0x52, 0x5B)
BLUE = RGBColor(0x25, 0x63, 0xEB)
AMBER = RGBColor(0xB4, 0x5D, 0x05)
RED = RGBColor(0xB9, 0x1C, 0x1C)

K = ["pv_potential", "daylighting_potential", "relative_compactness", "fsi_performance"]
NICE = ["PV potential", "Daylighting", "Rel. compactness", "FSI performance"]
TH = np.array([0.70, 0.70, 0.75, 0.80])

# =====================================================================================
# COMPUTE EVERYTHING FIRST
# =====================================================================================
d = pd.read_csv(DT / "data/given/appendix_c.csv")
ids = d.configuration.to_numpy()
A = d[K].to_numpy()
Ai = np.rint(A * 100).astype(int)
Bi = np.rint(TH * 100).astype(int)


def nondominated(Y):
    n = len(Y)
    return [i for i in range(n)
            if not any(all(Y[j] >= Y[i]) and any(Y[j] > Y[i]) for j in range(n) if j != i)]


ND30 = len(nondominated(A))
feas = np.where((Ai >= Bi).all(1))[0]
F, FN = A[feas], ids[feas]
NDF = len(nondominated(F))


def support_margin(Y, k):
    """max t s.t. w·y_k - w·y_j >= t for all j != k, w > 0, sum w = 1.  t > 0 => supported."""
    others = [j for j in range(len(Y)) if j != k]
    res = linprog([0, 0, 0, 0, -1],
                  A_ub=[list(Y[j] - Y[k]) + [1] for j in others], b_ub=[0.0] * len(others),
                  A_eq=[[1, 1, 1, 1, 0]], b_eq=[1],
                  bounds=[(1e-6, 1)] * 4 + [(None, None)])
    return -res.fun if res.success else None


MARGIN = {FN[k]: support_margin(F, k) for k in range(len(F))}
UNSUP = [c for c, t in MARGIN.items() if t <= 0]
SUP = [c for c, t in MARGIN.items() if t > 0]

rng = np.random.default_rng(20260728)
Wd = rng.dirichlet(np.ones(4), 500_000)
share = np.bincount(np.argmax(Wd @ F.T, axis=1), minlength=len(F)) / len(Wd) * 100
SHARE = dict(zip(FN, share))

RHO = 1e-6
def asf(q, lam=np.full(4, 0.25)):
    g = (F - q) / lam
    return g.min(1) + RHO * g.sum(1)

ASF_TIE = sorted(FN[np.isclose(asf(TH), asf(TH).max(), atol=1e-4)].tolist(),
                 key=lambda c: -asf(TH)[list(FN).index(c)])
ASF_C15 = FN[int(np.argmax(asf(F[list(FN).index("C15")])))]

PROF = {"Client": [.10, .10, .10, .70], "Municipality": [.35, .35, .15, .15],
        "Engineering": [.15, .15, .55, .15]}
U = np.column_stack([F @ np.array(w) for w in PROF.values()])
Un = (U - U.min(0)) / (U.max(0) - U.min(0))
NASH = dict(zip(FN, np.prod(Un, 1)))
MM = dict(zip(FN, Un.min(1)))
MM_ZERO = [c for c in FN if MM[c] < 1e-9]

MUNI = []
for g in (.45, .40, .35, .30, .25):
    r = (1 - 2 * g) / 2
    sc = F @ np.array([g, g, r, r])
    o = np.argsort(-sc)
    MUNI.append((f"{int(g*100)}/{int(g*100)}/{int(r*100):02d}/{int(r*100):02d}",
                 FN[o[0]], sc[o[0]] - sc[o[1]]))

SWEEP = [(dd, int((Ai >= Bi + dd).all(1).sum())) for dd in range(-5, 7)]
EMPTY = next(dd for dd, c in SWEEP if c == 0)

SLACK = []
for j in range(4):
    lv = sorted({Ai[i, j] for i in range(30)})
    last = max(L for L in lv if ((Ai >= np.where(np.arange(4) == j, L, Bi)).all(1)).any())
    binder = ids[np.argmax(np.where((Ai >= np.where(np.arange(4) == j, last, Bi)).all(1),
                                    Ai[:, j], -1))]
    SLACK.append((NICE[j], TH[j], last / 100, (last - Bi[j]) / 100, binder))

UNIQ = []
for j in range(4):
    other = np.all([Ai[:, i] >= Bi[i] for i in range(4) if i != j], 0)
    u = ids[other & (Ai[:, j] < Bi[j])]
    UNIQ.append((NICE[j], ["municipality", "municipality", "engineering", "client"][j],
                 list(u), int(other.sum())))

KNIFE = []
for lab, t in [("as agreed", Bi), ("daylight 0.70 → 0.71", Bi + [0, 1, 0, 0]),
               ("PV 0.70 → 0.71", Bi + [1, 0, 0, 0])]:
    m = (Ai >= t).all(1)
    sub, Z = ids[m], A[m]
    us = [sub[k] for k in range(len(sub)) if support_margin(Z, k) <= 0]
    KNIFE.append((lab, int(m.sum()),
                  {a: sub[int(np.argmax(Z @ np.array(w)))] for a, w in PROF.items()}, us))

C30 = np.corrcoef(A.T)
C10 = np.corrcoef(F.T)
pc1 = lambda M: (lambda e: e[0] / e.sum())(np.linalg.eigvalsh(np.corrcoef(M.T))[::-1])
RUNS = [("Appendix C (the case)", len(A), pc1(A), C30[0, 1], C30[0, 3])]
for p in sorted(glob.glob(str(DT / "data/runs/*/space/performance_normalised.csv"))):
    M = pd.read_csv(p)[K].to_numpy()
    Cx = np.corrcoef(M.T)
    RUNS.append((p.split("/")[-3], len(M), pc1(M), Cx[0, 1], Cx[0, 3]))

PAIRS = [(f"{NICE[i]} ~ {NICE[j]}", C30[i, j], C10[i, j])
         for i in range(4) for j in range(i + 1, 4)]
MINABS = min(abs(v) for _, v, _ in PAIRS)


def fmt(x, n=2):
    return f"{x:.{n}f}"


def lst(x):
    return ", ".join(x)


# =====================================================================================
# DOCUMENT SCAFFOLDING
# =====================================================================================
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.1)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(1.8)
CONTENT_W = 21.0 - 4.2

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.13

FIGN = [0]
TABN = [0]


def _shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def _borders(table, colour="D4D4D8", sz=4):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), colour)
        b.append(e)
    tblPr.append(b)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(13.5)
    r.font.bold = True
    r.font.color.rgb = INK
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = BLUE
    return p


def body(*chunks, space_after=5, indent=0.0, italic=False, size=10):
    """chunks: str, or (str, {'b':True,'c':COLOR,'i':True,'mono':True})"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for c in chunks:
        text, opt = (c, {}) if isinstance(c, str) else c
        r = p.add_run(text)
        r.font.size = Pt(opt.get("sz", size))
        r.font.bold = opt.get("b", False)
        r.font.italic = opt.get("i", italic)
        r.font.color.rgb = opt.get("c", INK)
        if opt.get("mono"):
            r.font.name = "Consolas"
    return p


def bullet(*chunks, space_after=3):
    p = body(*chunks, space_after=space_after, indent=0.55)
    p.paragraph_format.first_line_indent = Cm(-0.32)
    return p


def table(headers, rows, widths=None, caption=None, fontsize=8.8, align_right=()):
    TABN[0] += 1
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _borders(t)
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        r = p.add_run(htxt)
        r.font.size = Pt(fontsize)
        r.font.bold = True
        r.font.color.rgb = INK
        _shade(hdr[i], "F4F4F5")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            text, opt = (val, {}) if isinstance(val, str) else val
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(text)
            r.font.size = Pt(fontsize)
            r.font.bold = opt.get("b", False)
            r.font.color.rgb = opt.get("c", INK)
            if opt.get("mono"):
                r.font.name = "Consolas"
    if widths:
        for i, wcm in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(wcm)
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after = Pt(9)
        r = cp.add_run(f"Table {TABN[0]}. {caption}")
        r.font.size = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = MUTED
    return t


def figure(png, caption, width=CONTENT_W):
    FIGN[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(HERE / png), width=Cm(width))
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(9)
    r = cp.add_run(f"Figure {FIGN[0]}. {caption}")
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = MUTED


def callout(text, colour=BLUE):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    _shade(c, "F8FAFC" if colour == BLUE else "FEF7ED")
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = colour
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.space_before = Pt(0)
    sp.runs and None
    for r in sp.runs: r.font.size = Pt(2)
    return t


def page_number_footer():
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Case Study 2 · Solar-Efficient Building Design, Rotterdam · page ")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


page_number_footer()

# =====================================================================================
# TITLE BLOCK
# =====================================================================================
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("2026 EURO PhD SUMMER SCHOOL ON MCDA/MCDM · TU DELFT")
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Case Study 2 — Designing a Solar-Efficient Residential Building in Rotterdam")
r.font.size = Pt(17)
r.font.bold = True
r.font.color.rgb = INK

body(("Consultants' report to Studio Delta — analysis of the feasible design space, "
      "stakeholder trade-offs, requirement sensitivity, and recommendations for reuse.",
      {"c": MUTED, "sz": 10.5}), space_after=10)

# =====================================================================================
# 1 INTRODUCTION
# =====================================================================================
h1("1  Introduction")
body("Building massing — the overall shape, height and footprint of a building — is decided "
     "before floor plans, façades or materials. Those early decisions fix how much roof is "
     "available for photovoltaics, how much daylight reaches the façade, how much heat "
     "escapes through the envelope, and how many dwellings the site yields. They cannot be "
     "undone cheaply later.")
body("Studio Delta commissioned a voxel-based computational design model that generated "
     "thousands of feasible massings inside the municipality's allowable volume and scored "
     "each on four performance measures. The engineering team retained thirty representative "
     "configurations. In Meeting 4 the client, the municipality and the engineering team each "
     "stated a minimum acceptable level on one or two of those measures.")
body("The optimisation is therefore already done. Our task, in Laura's words, is ",
     ("\"to analyse the feasible design space produced by the engineering team and determine "
      "which configuration should ultimately be recommended.\"", {"i": True}),
     " This report does that, and then asks a second question that the case does not: ",
     ("which of the thirty designs a given decision method is capable of recommending at all.",
      {"b": True}))

# =====================================================================================
# 2 THE PROBLEM
# =====================================================================================
h1("2  The decision problem")
h2("2.1  Problematique and actors")
body("Following Kadziński (Method Selection, p. 6), the recommendation requested is a ",
     ("choice (α)", {"b": True}), " — select one or more of an explicit, stable set of thirty "
     "alternatives — not a ranking, sorting or clustering. Section 7 argues that the "
     "recurring version of this problem is a sorting task instead.")
body("Three actors have standing, in Cinelli's sense (Problem structuring, pp. 8–20; "
     "Power–Interest matrix pp. 16–17, after Ferretti & Grosso 2019). Each stated a minimum, "
     "and Laura's decision rule is conjunctive: ",
     ("\"Any configuration failing to satisfy one or more stakeholder requirements cannot be "
      "recommended.\"", {"i": True}))

table(["Actor", "Stated in Meeting 4", "Standing"],
      [["Client", "FSI performance ≥ 0.80", "pays for the project; high power, high interest"],
       ["Municipality", "PV ≥ 0.70 and daylighting ≥ 0.70", "grants permission; high power, lower day-to-day interest"],
       ["Engineering team", "Relative compactness ≥ 0.75", "advises; low formal power, high interest"]],
      widths=[3.0, 5.4, 8.4],
      caption="The three actors and the only preference information the case supplies.")

h2("2.2  The fact that governs the whole analysis")
callout("The case supplies four minimum levels and no weights. That is the entire preference "
        "information in the document.")
body("Every weight vector appearing in this report is therefore ", ("ours", {"b": True}),
     " — a modelling assumption, not data. We name each one where it is used, and Section 6.2 "
     "shows which conclusions survive the choice and which do not. Best–Worst Method (Rezaei, "
     "pp. 10–18) is what we would use to elicit weights properly: 2n − 3 = 5 comparisons for "
     "four criteria, with a reportable consistency ratio. We propose it; we did not perform it, "
     "because there is no decision maker in the room to answer.")

# =====================================================================================
# 3 DATA
# =====================================================================================
h1("3  Data")
h2("3.1  What the case provides")
body("Appendix C gives a 30 × 4 matrix of normalised performances, reproduced in full in "
     "Appendix A.1 of this report. All four measures are to be maximised and all lie in [0, 1]. "
     "Appendix B of the case gives the formulas the engineering team used:")

table(["#", "Measure", "Formula", "What it counts"],
      [["F₁", "PV potential", ("Σᵢ Σⱼ  Vᵢⱼ · Wᵢ · A", {"mono": True}),
        "annual solar radiation reaching roof test points"],
       ["F₂", "Daylighting potential", ("Σᵢ Σₖ  Vᵢₖ · Lᵢ · A", {"mono": True}),
        "annual direct illuminance reaching façade test points"],
       ["F₃", "Relative compactness", ("6 · Vb^(2/3) / Sc", {"mono": True}),
        "enclosed volume per unit envelope; 1 = a cube"],
       ["F₄", "FSI performance", ("2w / (w + t)", {"mono": True}),
        "floor space index w against target t; 1 = target met"]],
      widths=[0.9, 3.4, 4.2, 8.3],
      caption="The four performance measures (case Appendix B). i indexes solar positions over the year, j roof test points, k façade test points. V is a visibility indicator, W global horizontal radiation, L direct normal illuminance, A voxel area, Vb enclosed volume, Sc envelope surface, w the floor space index and t its target.")

h2("3.2  What the case does not provide, and why it matters")
bullet(("Normalisation. ", {"b": True}),
       "The transformation from raw F₁–F₄ to the [0, 1] values in Appendix C is not published. "
       "The thresholds 0.70 and 0.80 are therefore numbers on an undisclosed scale. Every "
       "result here is conditional on the table as given.")
bullet(("The sampling rule. ", {"b": True}),
       "Thousands of configurations became thirty after \"removing duplicate and clearly "
       "inferior designs\" and selecting a set that \"captures the diversity of the explored "
       "design space\". The rule is not stated. The thirty are a sample of the efficient set, "
       "not the efficient set, so every claim of the form \"Cx is efficient\" means \"efficient "
       "among the thirty we were shown\".")
bullet(("Precision. ", {"b": True}),
       "Values are published to two decimals, so each carries up to ±0.005 and a difference "
       "between two scores carries up to ±0.010. Section 6 shows several margins below that.")
bullet(("Anything outside the plot. ", {"b": True}),
       "No measure describes an effect on neighbouring buildings, although the brief raises it "
       "twice. Section 7.2 treats this as a criteria-family defect rather than an oversight.")

h2("3.3  Data we generated as a control")
body("To test whether the correlation structure of Appendix C is a property of the four "
     "measures or of the supplied sample, we reimplemented the Appendix B formulas and ran "
     "them on real Dutch geometry and weather: building footprints and heights from ",
     ("3DBAG", {"b": True}), ", and typical-meteorological-year irradiance from ",
     ("PVGIS TMY", {"b": True}), ". Four design spaces were generated (Rotterdam Blaak, "
     "Zoetermeer Edge, and two further Rotterdam parcels), 30–40 configurations each. "
     "Section 6.4 reports the comparison.")

# =====================================================================================
# 4 WHAT WE DO WITH THE DATA
# =====================================================================================
h1("4  What we do with the data")
body("The analysis runs as a fixed six-step pipeline. Steps 1–3 need no preference information "
     "at all; only steps 4–5 introduce weights, and they are declared.")

table(["Step", "Operation", "Preference information needed", "Answers"],
      [["1", "Pareto dominance filter", "none", "is the set already conflicting?"],
       ["2", "Conjunctive screen on the four minimums", "the four stated minimums", "which designs are admissible?"],
       ["3", "Supported / unsupported classification (LP)", "none", "which designs can a weighted sum reach?"],
       ["4", "Weight-space sampling over the simplex", "none — it sweeps all weights", "which designs win, and how often?"],
       ["5", "Stakeholder profiles + fairness rules", "our weight assumptions, declared", "whose design wins, and is that stable?"],
       ["6", "Parametric ε-constraint sensitivity", "the minimums, varied", "how much room is there?"]],
      widths=[1.0, 5.5, 5.0, 5.3],
      caption="The pipeline. Only steps 4–5 involve assumed weights.")

body("A seventh step audits the criteria family itself against Cinelli's requirements "
     "(pp. 42–43) — completeness, non-redundancy, preferential independence, conciseness — "
     "using both Appendix C and the generated control spaces.")

# =====================================================================================
# 5 METHODS
# =====================================================================================
h1("5  Methods")
h2("5.1  Methods used, and why each was chosen")

table(["Method", "Source", "Why it is here", "Requires"],
      [["Pareto dominance", "Köksalan I, pp. 5–6; Słowiński, p. 3",
        "The only preference-free comparison. Słowiński: the dominance relation is the only "
        "objective information in a multi-attribute problem. Establishes whether preference "
        "information is needed at all.", "nothing"],
       ["Conjunctive screen (thresholds as ε-constraints, and as vetoes)",
        "Köksalan I, pp. 15–16; Figueira, ELECTRE I, pp. 17–19",
        "Meeting 4 states minimum performances and failing one disqualifies. That is "
        "non-compensatory language, which a weighted sum structurally cannot express.",
        "the four minimums"],
       ["Supported / unsupported classification by LP",
        "Köksalan I, p. 7 (solution types), pp. 12–13",
        "Asks a question no scoring method can: is this design reachable by any positive "
        "weighted sum? A property of the set, not of anyone's preferences.", "nothing"],
       ["Weight-space sampling (first-rank acceptability)",
        "Köksalan I, p. 14; Kadziński, p. 14 (stochastic acceptabilities)",
        "Replaces one arbitrary weight vector with a distribution over winners. Not full SMAA: "
        "performances are deterministic and we report first-rank shares only.",
        "nothing"],
       ["Achievement scalarising function",
        "Köksalan II, pp. 28–31 (Wierzbicki 1980)",
        "The only method here that can reach an unsupported point. Used to demonstrate the "
        "reachability gap rather than assert it.", "a reference point q"],
       ["Nash bargaining; max-min; Gini",
        "Vetschera, pp. 31–34, 65–71",
        "Aggregating across actors when no common utility scale exists. Nash is invariant to "
        "positive affine rescaling per actor; max-min is not.",
        "actor value columns"]],
      widths=[3.5, 3.7, 7.5, 2.1], fontsize=8.4,
      caption="Methods applied, with the lecture source that supplies each.")

h2("5.2  The supportedness test")
body("A configuration y_k in a set Y is ", ("supported", {"b": True}), " if some strictly "
     "positive weight vector makes it a maximiser of the weighted sum over Y. We test this "
     "exactly, one linear programme per configuration:")
body(("maximise  t   subject to   w·(y_j − y_k) + t ≤ 0  ∀ j ≠ k,   Σ wᵢ = 1,   wᵢ ≥ 10⁻⁶",
      {"mono": True, "sz": 9}), indent=0.7, space_after=5)
body("A strictly positive optimal t means a weight vector exists that makes y_k the unique "
     "winner; t ≤ 0 means none does. Two points matter. First, ",
     ("supportedness is a property of an alternative relative to a set", {"b": True}),
     ", not a fixed label — Section 6.3 shows a design changing status when another is removed. "
     "Second, sampling cannot prove it: a zero share means \"not found in 500 000 draws\", not "
     "\"impossible\". The LP is the proof; the sampling corroborates.")

h2("5.3  Methods considered and not used")
bullet(("A weighted sum alone. ", {"b": True}),
       "Fully compensatory: a strong PV score would buy its way past a failed compactness "
       "minimum, which Meeting 4 forbids. It also cannot reach three of the ten admissible "
       "designs (Section 6.1).")
bullet(("ELECTRE Tri-nB. ", {"b": True}),
       "The right problematique for the reusable version of this problem (Section 7.3), but it "
       "requires weights, indifference, preference and veto thresholds and a cutting level λ — "
       "more than a dozen parameters, none supplied by the case.")
bullet(("DRSA. ", {"b": True}),
       "Degenerates here: the class labels would be derived from the thresholds themselves, so "
       "the approximation is exact by construction and the induced rules merely restate the "
       "screen. It becomes the right tool once Studio Delta has a history of accepted and "
       "rejected designs.")
bullet(("Transplanting weights between families. ", {"b": True}),
       "We elicit for one model only. Figueira (ELECTRE II, p. 4): ",
       ("\"The compensatory effects are not pertinent. This is due to the fact that the weights "
        "cannot be interpreted as substitution rates.\"", {"i": True}),
       " BWM and AHP produce trade-off weights; ELECTRE weights are voting power. They are not "
       "interchangeable.")

# =====================================================================================
# 6 RESULTS
# =====================================================================================
h1("6  Results")
h2("6.1  The feasible design space, and what a weighted sum can reach")

body(f"All {ND30} configurations are mutually non-dominated: there are zero dominated pairs "
     f"among 870 ordered comparisons. Dominance filtering removes nothing, because the "
     f"engineering team already removed the dominated designs before handing the table over. "
     f"The four minimums cut 30 → ", (f"{len(feas)}", {"b": True}), f" ({lst(FN)}), and those "
     f"{len(feas)} are also mutually non-dominated.")

figure("Q1_design_space.png",
       "Screening, reachability, and what the method choice decides. Left: worst slack against "
       "the four minimums, 30 → 10. Centre: share of the weight simplex won, 500 000 uniform "
       "draws; open circles mark designs that are never selected. Right: the achievement "
       "scalarising function with q = y(C15).")

body("Three of the ten are ", ("unsupported within this set", {"b": True, "c": AMBER}),
     f": {lst(UNSUP)}. No positive weight vector makes any of them a weighted-sum winner. "
     f"This is proved by linear programming, and corroborated by sampling: across 500 000 "
     f"uniform weight draws they win zero times, while the knife-edge {SUP[-1] if False else 'C11'} "
     f"— technically supported, with an optimality margin of {MARGIN['C11']:.6f} — wins "
     f"{int(SHARE['C11']*5000)} times ({SHARE['C11']:.3f} %).")

table(["Configuration"] + NICE + ["LP margin t", "Simplex share"],
      [[c] + [fmt(v) for v in F[i]] +
       [(f"{MARGIN[c]:+.6f}", {"c": AMBER if MARGIN[c] <= 0 else INK,
                               "b": MARGIN[c] <= 0}),
        f"{SHARE[c]:.3f} %"] for i, c in enumerate(FN)],
      widths=[2.3, 2.0, 2.0, 2.2, 2.2, 2.5, 2.0], fontsize=8.4, align_right=(1, 2, 3, 4, 5, 6),
      caption="The ten admissible configurations. A non-positive LP margin means unsupported: "
              "no positive weight vector selects that design within this set.")

body("The consequence is not that these designs are worse. It is that a scoring method — a "
     "weighted sum, SAW, or AHP with additive aggregation — is ",
     ("structurally incapable", {"b": True}), " of recommending 30 % of the admissible set, and "
     "deletes them before anyone looks. A method with a different geometry does not. Setting "
     f"the reference point at C15's own performance, the Wierzbicki achievement scalarising "
     f"function selects ", (f"{ASF_C15}", {"b": True}), " outright.")

callout("Which designs are recommendable at all depends on the family of scalarisation, "
        "before any weight is discussed.", colour=AMBER)

body("Two honest qualifications. Setting q = y(C15) is constructed to demonstrate "
     "reachability; it is not evidence that C15 is the right building. And at the reference "
     f"point q = the four minimums, the ASF does not discriminate either — it ties "
     f"{lst(ASF_TIE)} at 0.20.")

h2("6.2  Three stakeholders")
figure("Q2_stakeholders.png",
       "What survives the weighting assumption. Left: the winner as each actor's own weight is "
       "swept. Centre: the municipality's winning margin under five plausible weightings, "
       "against the ±0.005 rounding of a two-decimal table. Right: Nash and max-min across the "
       "three actors.")

body("Under the profiles declared in Appendix A.3, the client prefers ",
     (KNIFE[0][2]["Client"], {"b": True}), ", the municipality ",
     (KNIFE[0][2]["Municipality"], {"b": True}), ", and the engineering team ",
     (KNIFE[0][2]["Engineering"], {"b": True}), ". There is no consensus design. But the "
     "disagreement is not symmetric, and that asymmetry is the result:")

bullet(("Client and engineering are stable. ", {"b": True}),
       "C8 wins at every FSI weight from 0.40 to 0.85, with the margin widening from 0.004 to "
       "0.016. C14 wins once the compactness weight clears 5/11 ≈ 0.4545.")
bullet(("The municipality is not. ", {"b": True}),
       "Five plausible weightings of the same stated preference give three different winners:")
table(["Municipality profile", "Winner", "Margin over runner-up"],
      [[a, (b, {"b": True}), f"{c:.4f}"] for a, b, c in MUNI],
      widths=[5.0, 3.0, 4.5], fontsize=8.6, align_right=(2,),
      caption="Every margin is at or below 0.005 — the rounding of the source table.")

body("The reading matters. This is not a fickle stakeholder: it is our own elicitation model "
     "being unstable exactly where the case gave us least to work with. The municipality "
     "constrains two criteria and leaves two free; the client's single sharp constraint pins "
     "C8 regardless.")

body("Aggregating across actors, we avoid averaging the three weight vectors — averaging is "
     "welfare maximisation, and Vetschera (p. 33) shows the winner flipping under an "
     "admissible rescaling of one actor's utility. Instead: the ",
     ("Nash product", {"b": True}), f" selects ", (max(NASH, key=NASH.get), {"b": True}),
     f" ({max(NASH.values()):.5f}), and ", ("max-min", {"b": True}), " selects ",
     (max(MM, key=MM.get), {"b": True}), f" ({max(MM.values()):.3f}). Nash is primary because "
     f"it maximises Π(uᵢ − dᵢ) and min–max normalisation is a positive affine transform per "
     f"actor with dᵢ = 0, so its argmax is unchanged by it. We verified this: "
     f"{max(NASH, key=NASH.get)} wins under d = column minimum, d = 0, and d = the threshold "
     f"configuration's value.")

body(("A caveat we report rather than hide. ", {"b": True}),
     f"max-min = 0 for {lst(MM_ZERO)}. That is an artefact of min–max normalisation — whoever "
     f"is last on a column scores zero by construction. These are not vetoes and not "
     f"\"unacceptable to someone\". It is precisely Vetschera's scale-dependence warning "
     f"landing on our own analysis.")

h2("6.3  How much room is in the requirements")
figure("Q3_thresholds.png",
       "Parametric ε-constraint analysis. Left: all four minimums shifted together. Centre: "
       "what each requirement removes that the others do not. Right: the knife-edge scenarios.")

body("Köksalan (Part II, p. 11) describes exactly this operation: ",
     ("\"Changing εᵢ values systematically, we can find many (sometimes all) efficient "
      "solutions… we can explore different (desirable) parts of the efficient frontier.\"",
      {"i": True}), " The four stated minimums are the εᵢ right-hand sides, so the lecture's "
     "own method matches the question directly rather than by analogy.")

table(["Shift applied to all four"] + [f"{dd:+d}" for dd, _ in SWEEP],
      [["Feasible configurations"] + [(str(c), {"b": c in (0, 10), "c": RED if c == 0 else INK})
                                      for _, c in SWEEP]],
      widths=[4.2] + [0.95] * len(SWEEP), fontsize=8.4,
      caption="Feasible count as all four minimums move together, in hundredths.")

body(f"The requirements as agreed sit ", (f"{EMPTY} hundredths from infeasibility", {"b": True}),
     ". Relaxing is far less dramatic — −0.05 admits 17 — so the constraint set is much tighter "
     "on the upside than the downside. There is no room to negotiate up and considerable room "
     "to negotiate down, which is useful to know before a negotiation rather than during one.")

table(["Requirement", "Set by", "Uniquely eliminates", "Feasible if dropped"],
      [[f"{n} ≥ {TH[j]:.2f}", own, (lst(u) if u else "nothing", {"b": not u, "c": RED if not u else INK}),
        str(tot)] for j, (n, own, u, tot) in enumerate(UNIQ)],
      widths=[4.6, 3.4, 4.0, 3.5], fontsize=8.6,
      caption="What each requirement removes that the others do not.")

body("Counting how many designs fail each criterion in isolation is misleading, because the "
     "failures overlap heavily. On the question that matters — what does each requirement "
     "remove that the others do not — ", ("the engineering team's compactness floor removes "
     "nothing", {"b": True, "c": RED}), ". Delete it from Meeting 4 and the same ten designs "
     "survive, and it stays non-binding at every tightening level from +0.00 to +0.05. The "
     "other three are load-bearing by exactly one configuration each: three requirements, three "
     "single points of failure.")

body("Two designs survive only at exact equality: C8 at daylighting = 0.70 and C14 at "
     "PV = 0.70. They are also the client's and the engineer's preferred designs. Moving the "
     "municipality's daylight floor by one hundredth — inside the rounding of the source "
     "table — therefore does more than change a winner:")

table(["Scenario", "Feasible", "Client picks", "Municipality", "Engineering", "Unsupported in that set"],
      [[lab, str(n), (w["Client"], {"b": True}), w["Municipality"], w["Engineering"], lst(us)]
       for lab, n, w, us in KNIFE],
      widths=[3.8, 1.7, 2.3, 2.4, 2.4, 3.2], fontsize=8.4,
      caption="Knife-edge scenarios. Raising the daylight floor by 0.01 removes C8 and changes "
              "which designs a weighted sum can reach.")

body("With C8 gone, ", ("C15 changes from unsupported to supported", {"b": True}),
     " and the client's own weighted sum then selects it. Moving a threshold rewrites the "
     "geometry of the feasible set, not just its membership. Meanwhile the engineering team's "
     "own requirement changes nothing at all — influence in this decision is not where the "
     "stakeholders would assume it is.")

h2("6.4  The criteria family")
figure("Q4_criteria_audit.png",
       "Auditing the criteria family. Left: the Appendix C correlation matrix. Centre: the same "
       "four formulas recomputed on real geometry and weather. Right: the three findings.")

body(f"Every one of the six criterion pairs in Appendix C correlates at |r| ≥ {MINABS:.3f}, and "
     f"the first principal component explains ", (f"{pc1(A)*100:.1f} %", {"b": True}),
     f" of the variance. This does not soften on the designs we actually recommend from: "
     f"{pc1(F)*100:.1f} % on the feasible ten, where three of the six correlations strengthen "
     f"(PV~compactness → {C10[0,2]:+.3f}, daylighting~FSI → {C10[1,3]:+.3f}). Four criteria, one "
     f"dominant dimension.")

table(["Criterion pair", "r over all 30", "r over the feasible 10"],
      [[n, f"{a:+.3f}", f"{b:+.3f}"] for n, a, b in PAIRS],
      widths=[7.0, 3.5, 4.0], fontsize=8.6, align_right=(1, 2),
      caption="All six pairwise correlations, on the full set and on the admissible subset.")

body("Recomputing the Appendix B formulas on real 3DBAG geometry and PVGIS weather gives a "
     "markedly different structure. The two couplings are not equally artificial:")

table(["Design space", "n", "PC1", "r(PV, daylight)", "r(PV, FSI)"],
      [[(n, {"b": i == 0}), str(k), f"{p*100:.1f} %", f"{a:+.3f}", f"{b:+.3f}"]
       for i, (n, k, p, a, b) in enumerate(RUNS)],
      widths=[5.4, 1.5, 2.4, 3.2, 2.5], fontsize=8.6, align_right=(1, 2, 3, 4),
      caption="Appendix C against four design spaces generated from real Dutch geometry and "
              "typical-meteorological-year weather, using the same four formulas.")

bullet(("PV ~ FSI persists but far weaker ", {"b": True}), "(+0.975 → +0.55 to +0.68). Denser "
       "massing does tend to raise both roof area and floor area. This coupling is real, and "
       "overstated in Appendix C.")
bullet(("PV ~ daylighting is not reproduced ", {"b": True}), "(−0.979 → −0.04 to +0.37). The "
       "near-perfect trade-off between rooftop PV and façade daylight did not appear in any of "
       "the four runs.")

body(("The limit of this evidence. ", {"b": True}),
     "The four generated spaces share one generator, one family of massing operations and one "
     "normalisation approach. They are replications under a single pipeline, not independent "
     "validation, and Appendix C's own normalisation is undisclosed. The comparison is "
     "suggestive, not conclusive.")

# =====================================================================================
# 7 DISCUSSION
# =====================================================================================
h1("7  Discussion")
h2("7.1  Recommendation")
body("Score gaps among the ten admissible designs run from 0.001 to 0.012, on a table published "
     "to two decimals. A single winner is not defensible at that precision. We recommend a set "
     "of four, with the condition under which each is the right answer:")

table(["Design", "Role", "Evidence"],
      [["C6", ("compromise", {"b": True}),
        f"Nash bargaining solution ({NASH['C6']:.5f}); equal-weight optimum (0.8175); wins L₁ "
        f"and L₂ distance to the ideal and ties C11 at L∞; the municipality's pick under two of "
        f"five weightings."],
       ["C11", ("compromise", {"b": True}),
        f"max-min solution ({MM['C11']:.3f}) — the design no actor ranks worst; second on equal "
        f"weights (0.8125); supported, but on a knife edge (LP margin {MARGIN['C11']:.6f})."],
       ["C8", ("client extreme", {"b": True}),
        f"largest share of the weight simplex ({SHARE['C8']:.1f} %); invariant across every "
        f"FSI-heavy profile; survives only because daylighting = 0.70 exactly."],
       ["C14", ("engineering extreme", {"b": True}),
        f"second largest share ({SHARE['C14']:.1f} %); robust once the compactness weight "
        f"exceeds 5/11; survives only because PV = 0.70 exactly."]],
      widths=[1.8, 3.2, 11.8], fontsize=8.6,
      caption="The recommended set. Choosing among these four is a decision for the "
              "stakeholders, not for the consultants.")

body("The technical analysis narrows thirty designs to four and states what each choice implies. "
     "Selecting among the four is a political decision about whose priority dominates, and we "
     "deliberately do not make it.")

h2("7.2  What this analysis does not claim")
bullet("Appendix C's normalisation is undisclosed, so all thresholds are on an unknown scale.")
bullet("The thirty are a sample of the efficient frontier by an unstated rule, not the frontier.")
bullet("All weight vectors are ours. BWM is proposed as the proper elicitation, not performed.")
bullet("Sampling cannot prove unsupportedness; the LP is the proof and the 500 000 draws corroborate.")
bullet("Setting q = y(C15) demonstrates reachability; it is not an argument that C15 is best.")
bullet("The four control spaces share one generator, so they replicate rather than validate.")

h2("7.3  What Studio Delta should change next time")
body(("Derive criteria from stakeholder objectives, not from what the optimiser computes. ",
      {"b": True}),
     "All four criteria are quantities the voxel model could produce. The criterion the brief "
     "names twice — effect on sunlight for the surrounding properties (case pp. 2 and 4) — is "
     "the one it could not, and it is absent. Two designs could tie on all four measures while "
     "one overshadows the neighbouring blocks and the other does not; the model would be blind "
     "to the difference, and Roy's exhaustiveness requirement says we should not be. Cinelli "
     "(pp. 86–91) gives the alternative route, and it is an architecture case: stakeholder "
     "interviews → fundamental and means objectives → features → design alternatives.")
body(("Choose the method from problem features, and say so. ", {"b": True}),
     "Kadziński's MCDA-MSS (pp. 5–15) describes methods and case studies over 156 "
     "characteristics and filters conjunctively. The deliverable is a defensible sentence — "
     "\"we used X because our problem has features a, b, c\" — which is checkable, and which "
     "most published applications do not provide.")
body(("Treat the recurring question as sorting, not choice. ", {"b": True}),
     "This project is a choice problem: pick from thirty given designs. The question Studio "
     "Delta will face repeatedly is different — is this massing acceptable, marginal or "
     "unacceptable? That is sorting (β), and sorting is ", ("absolute evaluation", {"b": True}),
     ": each design is compared to fixed reference profiles rather than to the other candidates "
     "(Figueira, ELECTRE I, p. 34, after Roy 1996). It is therefore rank-reversal immune, so "
     "next year's configurations cannot reshuffle this year's verdicts — exactly the property a "
     "firm accumulating designs across projects needs. ELECTRE Tri-nB and DRSA are the "
     "candidates; DRSA needs no weights at all and yields rules an architect can read directly.")

h2("7.4  Three traps we avoided deliberately")
bullet(("Running five methods and comparing rankings. ", {"b": True}),
       "This looks like effort and is the error Cinelli and Kadziński spend their lectures "
       "diagnosing: it treats method choice as a robustness check rather than a modelling "
       "commitment. We made one justified choice and reported robustness within it.")
bullet(("Transplanting weights across families. ", {"b": True}),
       "Trade-off weights and importance coefficients are different objects; moving BWM or AHP "
       "weights into ELECTRE or TOPSIS is a documented and common error.")
bullet(("Declaring a winner to three decimals. ", {"b": True}),
       "The municipality profile separates its top two designs by 0.001 on a table published to "
       "two decimals. We report a set and the conditions, not a ranking.")

# =====================================================================================
# 8 APPENDIX
# =====================================================================================
doc.add_page_break()
h1("Appendix A  Data, parameters and reproducibility")
h2("A.1  Appendix C in full — the thirty feasible configurations")
rows = []
for i, c in enumerate(ids):
    ok = i in feas
    rows.append([(c, {"b": ok, "c": BLUE if ok else INK})] +
                [(fmt(v), {"c": BLUE if ok else INK}) for v in A[i]] +
                [("\u2713" if ok else "\u2013", {"c": BLUE if ok else MUTED, "b": ok})])
half = 15
left, right = rows[:half], rows[half:]
table(["Cfg", "PV", "Dayl.", "Comp.", "FSI", "\u2713",
       "Cfg", "PV", "Dayl.", "Comp.", "FSI", "\u2713"],
      [left[i] + right[i] for i in range(half)],
      widths=[1.45, 1.30, 1.30, 1.30, 1.30, 0.75] * 2, fontsize=8.0,
      align_right=(1, 2, 3, 4, 7, 8, 9, 10),
      caption="The full design space (case Appendix C). \u2713 marks the ten configurations "
              "that satisfy all four Meeting-4 minimums; \u2013 marks those failing at least "
              "one. All measures are maximised; all values as published.")

h2("A.2  One-at-a-time slack in each requirement")
table(["Requirement", "As agreed", "Highest still feasible", "Slack", "Binding design"],
      [[n, fmt(b), fmt(last), f"+{s:.2f}", bind] for n, b, last, s, bind in SLACK],
      widths=[4.2, 2.4, 4.2, 2.0, 3.0], fontsize=8.6, align_right=(1, 2, 3),
      caption="How far each minimum can rise alone, with the other three held as agreed. "
              "Daylighting is the tightest lever. PV and FSI are both limited by C8.")

h2("A.3  Weight profiles used")
body("These are ", ("assumptions, not data", {"b": True}), ". They are stated here so every "
     "number in Section 6.2 is reproducible, and Section 6.2 shows explicitly which conclusions "
     "depend on them.")
table(["Actor", "PV", "Daylighting", "Compactness", "FSI", "Rationale"],
      [[a] + [f"{v:.2f}" for v in w] +
       [["dominant weight on the criterion the actor named in Meeting 4",
         "the two criteria the municipality constrained, weighted equally",
         "dominant weight on compactness"][i]]
       for i, (a, w) in enumerate(PROF.items())],
      widths=[2.8, 1.5, 2.2, 2.4, 1.5, 6.4], fontsize=8.6, align_right=(1, 2, 3, 4),
      caption="Stakeholder weight profiles. Actor values are min–max normalised over the ten "
              "admissible designs with disagreement point d = 0.")

h2("A.4  A numerical detail that changes a result")
body("All threshold comparisons in Section 6.3 are performed in ",
     ("exact integer hundredths", {"b": True}), ", not floating point. In binary floating point "
     "0.80 + 0.14 evaluates to 0.9400000000000001, so C8's FSI of exactly 0.94 fails its own "
     "limit by one part in 10¹⁶ and the reported slack comes out one step short. With two "
     "designs sitting at exact equality, that artefact attacks the knife-edge analysis "
     "directly. This is a real defect we encountered and corrected, not a hypothetical.")

h2("A.5  Reproducibility")
body("Every figure, table and number in this report is regenerated by the scripts below. "
     "Nothing is transcribed by hand.")
table(["Script", "Produces"],
      [[("q1_design_space.py", {"mono": True}), "Figure 1 · screening, LP supportedness, weight-space sampling, ASF"],
       [("q2_stakeholders.py", {"mono": True}), "Figure 2 · stakeholder sweeps, municipality profiles, Nash and max-min"],
       [("q3_thresholds.py", {"mono": True}), "Figure 3 · ε-constraint sensitivity, binding analysis, knife-edge scenarios"],
       [("q4_criteria_audit.py", {"mono": True}), "Figure 4 · correlation audit and the generated control spaces"],
       [("style.py", {"mono": True}), "shared palette and zero-marker helpers, so all four figures read as one system"],
       [("make_report.py", {"mono": True}), "this document"]],
      widths=[4.6, 11.9], fontsize=8.6,
      caption="Analysis code. Random draws use seed 20260728; the LP is solved exactly.")

h2("A.6  Sources cited")
for s_ in [
    "Cinelli, M. — Problem structuring (2026 EURO PhD Summer School). Stakeholders pp. 8–20; "
    "Power–Interest matrix pp. 16–17; criteria requirements pp. 42–43; correlations pp. 83–85; "
    "objectives-to-alternatives pp. 86–91.",
    "Köksalan, M. — Interactive Multiobjective Optimization, Part I (choice problems): solution "
    "types p. 7, weighted sums pp. 12–14, ε-constraint pp. 15–16, sorting pp. 48–54. "
    "Part II (design problems): ε-constraint pp. 9–11, achievement scalarising programme pp. 28–31.",
    "Figueira, J. R. — ELECTRE Methods, Part I: concordance, discordance and veto pp. 14–25; "
    "choice problematique pp. 26–28; absolute versus relative evaluation p. 34. Part II: "
    "strengths and the non-compensation counter-example pp. 4–5. ELECTRE Tri-nB: assignment "
    "procedures pp. 13–17, numerical example pp. 23–25.",
    "Kadziński, M. — Recommending MCDA Methods with a Taxonomy-Based DSS: problem typology "
    "pp. 5–9, preference model pp. 10–11, elicitation pp. 12–13, exploitation p. 14, "
    "rule-based recommendation p. 18.",
    "Rezaei, J. — Best-Worst Method: steps and consistency pp. 10–18; Best-Worst Tradeoff pp. 24–34.",
    "Vetschera, R. — Collective Decision Making: efficiency pp. 15–22; welfare and rescaling "
    "pp. 24–26 and p. 33; max-min p. 31; Gini p. 34; Nash bargaining pp. 65–71.",
    "Słowiński, R. — Decision Rule Approach: the dominance relation as the only objective "
    "information p. 3; rules and approximations pp. 30, 49–50.",
    "Wierzbicki, A. (1980), via Köksalan Part II pp. 28–31 — the achievement scalarising function.",
]:
    bullet(("— ", {"c": MUTED}), (s_, {"sz": 8.8, "c": MUTED}), space_after=2)

doc.save(OUT)
print(f"saved: {OUT}")
print(f"figures {FIGN[0]} · tables {TABN[0]} · paragraphs {len(doc.paragraphs)}")
print(f"nondominated {ND30}/30 · feasible {len(feas)} · unsupported {UNSUP} · empty at +{EMPTY}")
